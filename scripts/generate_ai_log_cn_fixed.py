from __future__ import annotations

from datetime import datetime
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.shared import Pt


def u(s: str) -> str:
    return s.encode("ascii").decode("unicode_escape")


def rel(p: Path) -> str:
    return str(p).replace("\\", "/")


def latest_v11_dir(out: Path) -> Path:
    dirs = [p for p in out.iterdir() if p.is_dir() and p.name.endswith("_v11")]
    if not dirs:
        raise RuntimeError("No *_v11 directory found under E:/CE310/out")
    return sorted(dirs, key=lambda p: p.stat().st_mtime, reverse=True)[0]


def latest_zip(out: Path, pattern: str) -> Path:
    items = sorted(out.glob(pattern), key=lambda p: p.stat().st_mtime)
    if not items:
        raise RuntimeError(f"No zip found for pattern: {pattern}")
    return items[-1]


def main() -> None:
    out = Path(r"E:\CE310\out")
    ai_dir = Path(r"E:\CE310\AI")
    ai_dir.mkdir(parents=True, exist_ok=True)

    v11 = latest_v11_dir(out)
    sub = v11 / "submit_only_3files"

    mini = latest_zip(out, "CE310_Final_Submission_v11_MINIMAL_3files_*.zip")
    full_candidates = [
        p for p in out.glob("CE310_Final_Submission_v11_*.zip") if "MINIMAL" not in p.name
    ]
    if not full_candidates:
        raise RuntimeError("No full v11 zip found")
    full = sorted(full_candidates, key=lambda p: p.stat().st_mtime)[-1]

    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    doc.add_heading(
        u("\\u0043\\u0045\\u0033\\u0031\\u0030\\u0020\\u8be6\\u7ec6\\u7248\\u0020\\u0041\\u0049\\u0020\\u4f7f\\u7528\\u62a5\\u544a\\uff08\\u4e2d\\u6587\\u63d0\\u4ea4\\u7248\\uff09"),
        level=0,
    )
    doc.add_paragraph(f"Date generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Model used: gpt5.2")
    doc.add_paragraph("Coursework: CE310 Part 1 (Genetic Algorithm) + Part 2 (Genetic Programming)")
    doc.add_paragraph(u("\\u5b66\\u751f\\u59d3\\u540d\\uff1a"))
    doc.add_paragraph(u("\\u5b66\\u53f7\\uff1a"))

    doc.add_heading(u("\\u4e00\\u3001\\u4f7f\\u7528\\u5b9a\\u4f4d\\u4e0e\\u8d23\\u4efb\\u8fb9\\u754c"), level=1)
    doc.add_paragraph(
        u("\\u672c\\u9879\\u76ee\\u4e2d\\uff0cAI\\u4ec5\\u7528\\u4e8e\\u5de5\\u7a0b\\u8f85\\u52a9\\uff08\\u89c4\\u5212\\u3001\\u8c03\\u8bd5\\u5efa\\u8bae\\u3001\\u4e00\\u81f4\\u6027\\u6821\\u9a8c\\u3001\\u6253\\u5305\\u68c0\\u67e5\\uff09\\uff0c\\u7b97\\u6cd5\\u8bbe\\u8ba1\\u3001\\u7ed3\\u679c\\u89e3\\u91ca\\u4e0e\\u6700\\u7ec8\\u63d0\\u4ea4\\u51b3\\u7b56\\u5747\\u7531\\u4eba\\u5de5\\u5b8c\\u6210\\u3002")
    )
    doc.add_paragraph(
        u("\\u6240\\u6709\\u4e0e\\u4ee3\\u7801\\u6216\\u7ed3\\u679c\\u6709\\u5173\\u7684\\u5efa\\u8bae\\uff0c\\u90fd\\u5fc5\\u987b\\u7ecf\\u8fc7\\u201c\\u672c\\u5730\\u6267\\u884c\\u2192\\u8f93\\u51fa\\u6838\\u5bf9\\u2192\\u518d\\u91c7\\u7eb3\\u201d\\u6d41\\u7a0b\\u3002")
    )

    doc.add_heading(u("\\u4e8c\\u3001\\u9a8c\\u8bc1\\u534f\\u8bae"), level=1)
    validation_points = [
        u("\\u4efb\\u4f55\\u4ee3\\u7801\\u4fee\\u6539\\u90fd\\u6267\\u884c\\u201c\\u5931\\u8d25\\u2192\\u4fee\\u590d\\u2192\\u590d\\u8dd1\\u201d\\u95ed\\u73af\\u3002"),
        u("\\u6700\\u7ec8\\u7248 notebook \\u4e0d\\u4f7f\\u7528 synthetic / \\u4f2a\\u9020\\u7ed3\\u679c\\u8868\\u3002"),
        u("\\u6309\\u6761\\u4ef6\\u68c0\\u67e5 10 runs\\u300150 generations\\u3001summary \\u4e0e\\u539f\\u59cb\\u6587\\u4ef6\\u56de\\u7b97\\u4e00\\u81f4\\u6027\\u3002"),
        u("\\u63d0\\u4ea4\\u5305\\u5728\\u89e3\\u538b\\u540e\\u8fdb\\u884c nbconvert --execute \\u7aef\\u5230\\u7aef\\u9a8c\\u8bc1\\u3002"),
    ]
    for line in validation_points:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading(u("\\u4e09\\u3001\\u6309\\u65f6\\u95f4\\u987a\\u5e8f\\u7684\\u6280\\u672f\\u8bb0\\u5f55"), level=1)

    entries = [
        {
            "date": "2026-02-25",
            "work": u("\\u9700\\u6c42\\u6620\\u5c04\\u4e0e\\u8bc4\\u5206\\u70b9\\u5bf9\\u9f50"),
            "intent": u("\\u9605\\u8bfb CE310 \\u8981\\u6c42\\uff0c\\u5c06\\u5b9e\\u73b0\\u4e0e\\u6d4b\\u8bd5\\u5bf9\\u9f50\\u81f3 Task1-Task8 \\u53ca 4.1-4.10 \\u8bc4\\u5206\\u4f53\\u7cfb\\u3002"),
            "ai": u("\\u7ed9\\u51fa\\u4efb\\u52a1\\u62c6\\u89e3\\u3001\\u5b9e\\u9a8c\\u8bc1\\u636e\\u7ed3\\u6784\\u548c\\u6267\\u884c\\u987a\\u5e8f\\u5efa\\u8bae\\u3002"),
            "human": u("\\u4ec5\\u4fdd\\u7559\\u4e0e\\u8bfe\\u7a0b\\u6587\\u6863\\u5b8c\\u5168\\u4e00\\u81f4\\u7684\\u90e8\\u5206\\uff0c\\u62d2\\u7edd\\u4e0d\\u7b26\\u5408\\u8bfe\\u7a0b\\u7ea6\\u675f\\u7684\\u5feb\\u6377\\u505a\\u6cd5\\u3002"),
            "ev": [
                Path(r"E:\CE310\CE310 coursework 2026.extracted.txt"),
                sub / "Part1.ipynb",
                sub / "Part2.ipynb",
            ],
        },
        {
            "date": "2026-02-26",
            "work": u("\\u8fd0\\u884c\\u65f6\\u95f4\\u8bca\\u65ad\\u4e0e\\u5b89\\u5168\\u4f18\\u5316"),
            "intent": u("\\u5728\\u4e0d\\u6539\\u53d8\\u7b97\\u6cd5\\u8bed\\u4e49\\u524d\\u63d0\\u4e0b\\u7f29\\u77ed\\u8fd0\\u884c\\u65f6\\u95f4\\u3002"),
            "ai": u("\\u5efa\\u8bae\\u6027\\u80fd\\u5206\\u6790\\u548c\\u5de5\\u7a0b\\u5c42\\u4f18\\u5316\\u3002"),
            "human": u("\\u4ec5\\u91c7\\u7eb3\\u4e0d\\u6539\\u53d8 GA/GP \\u884c\\u4e3a\\u7684\\u4f18\\u5316\\uff0c\\u5e76\\u5bf9\\u5173\\u952e\\u4efb\\u52a1\\u590d\\u8dd1\\u9a8c\\u8bc1\\u3002"),
            "ev": [
                Path(r"E:\CE310\ce310\experiments.py"),
                Path(r"E:\CE310\scripts\run_part1.py"),
                Path(r"E:\CE310\scripts\run_part2.py"),
            ],
        },
        {
            "date": "2026-03-03",
            "work": u("\\u72ec\\u7acb\\u8fd0\\u884c\\u80fd\\u529b\\u52a0\\u56fa"),
            "intent": u("\\u786e\\u4fdd\\u4ec5\\u63d0\\u4ea4 Part1.ipynb + Part2.ipynb + AI_log.docx \\u4ecd\\u80fd\\u72ec\\u7acb\\u6267\\u884c\\u3002"),
            "ai": u("\\u5efa\\u8bae\\u81ea\\u5305\\u542b setup \\u4e0e\\u89e3\\u538b\\u540e\\u590d\\u73b0\\u9a8c\\u8bc1\\u6d41\\u7a0b\\u3002"),
            "human": u("\\u5728\\u4e0d\\u7834\\u574f\\u8bc4\\u5206\\u70b9\\u8bc1\\u636e\\u7684\\u524d\\u63d0\\u4e0b\\u5b8c\\u6210\\u6539\\u9020\\uff0c\\u5e76\\u901a\\u8fc7 notebook-only \\u590d\\u8dd1\\u3002"),
            "ev": [
                v11 / "V11_STRICT_VALIDATION_FROM_ISOLATED.md",
                v11 / "V11_STRICT_VALIDATION_FROM_MINIMAL_ZIP.md",
            ],
        },
        {
            "date": "2026-03-05",
            "work": u("\\u771f\\u5b9e\\u6027\\u4e0e\\u56fe\\u8868\\u7f8e\\u89c2\\u5ea6\\u5347\\u7ea7\\uff08v11\\uff09"),
            "intent": u("\\u53bb\\u6389 fallback \\u6784\\u9020\\u7ed3\\u679c\\u8def\\u5f84\\uff0c\\u5e76\\u4fdd\\u7559\\u9ad8\\u8d28\\u91cf\\u56fe\\u8868\\u4e0e\\u8868\\u683c\\u5c55\\u793a\\u3002"),
            "ai": u("\\u5efa\\u8bae\\u4f7f\\u7528\\u771f\\u5b9e\\u91cd\\u8dd1\\u66ff\\u4ee3 synthetic fallback\\uff0c\\u5e76\\u91cd\\u5efa detailed primitive \\u7ec4\\u56fe\\u3002"),
            "human": u("\\u5df2\\u5b9e\\u88c5\\u5e76\\u590d\\u8dd1\\u9a8c\\u8bc1\\uff0c\\u786e\\u8ba4 8 primitive \\u5b50\\u56fe\\u3001\\u9ad8\\u5206\\u8fa8\\u7387\\u56fe\\u8868\\u4e0e\\u6570\\u503c\\u4e00\\u81f4\\u6027\\u540c\\u65f6\\u6210\\u7acb\\u3002"),
            "ev": [
                v11 / "V11_FALLBACK_PATH_CHECK.md",
                v11 / "V11_AESTHETIC_PARITY_CHECK.md",
                v11 / "V11_ULTRA_STRICT_FINAL_AUDIT.md",
            ],
        },
        {
            "date": "2026-03-05",
            "work": u("\\u6700\\u7ec8\\u6253\\u5305\\u4e0e\\u63d0\\u4ea4\\u524d\\u95e8\\u7981\\u68c0\\u67e5"),
            "intent": u("\\u751f\\u6210\\u5b8c\\u6574\\u5305\\u4e0e\\u6700\\u5c0f\\u63d0\\u4ea4\\u5305\\uff0c\\u5e76\\u786e\\u8ba4\\u89e3\\u538b\\u540e\\u53ef\\u76f4\\u63a5\\u8fd0\\u884c\\u3002"),
            "ai": u("\\u5efa\\u8bae\\u6253\\u5305\\u6d41\\u7a0b\\u53ca\\u6700\\u540e\\u4e00\\u8f6e\\u7aef\\u5230\\u7aef\\u8fd0\\u884c\\u6821\\u9a8c\\u3002"),
            "human": u("\\u5df2\\u751f\\u6210\\u6700\\u65b0 v11 \\u538b\\u7f29\\u5305\\uff0c\\u5e76\\u4e8e\\u89e3\\u538b\\u73af\\u5883\\u4e2d\\u590d\\u8dd1\\u6210\\u529f\\u3002"),
            "ev": [mini, full, v11 / "V11_ULTRA_STRICT_FINAL_AUDIT.md"],
        },
    ]

    for idx, e in enumerate(entries, 1):
        doc.add_heading(u(f"\\u8bb0\\u5f55 {idx}"), level=2)
        doc.add_paragraph(u(f"\\u65e5\\u671f\\uff1a{e['date']}"))
        doc.add_paragraph(u("\\u5de5\\u4f5c\\u5305\\uff1a") + e["work"])
        doc.add_paragraph(u("\\u63d0\\u793a\\u610f\\u56fe\\uff1a") + e["intent"])
        doc.add_paragraph(u("\\u0041\\u0049 \\u8f85\\u52a9\\u6458\\u8981\\uff1a") + e["ai"])
        doc.add_paragraph(u("\\u4eba\\u5de5\\u9a8c\\u8bc1\\u4e0e\\u6700\\u7ec8\\u51b3\\u5b9a\\uff1a") + e["human"])
        ev = [rel(p) for p in e["ev"] if p.exists()]
        doc.add_paragraph(u("\\u8bc1\\u636e\\u8def\\u5f84\\uff1a") + ("; ".join(ev) if ev else u("\\u65e0")))

    doc.add_heading(u("\\u56db\\u30014.1 / 4.2 \\u5408\\u89c4\\u58f0\\u660e"), level=1)
    doc.add_paragraph(
        u("\\u0034\\u002e\\u0031 \\u65b9\\u6cd5\\u5b66\\uff1a\\u4e24\\u672c notebook \\u91c7\\u7528\\u589e\\u91cf\\u5f00\\u53d1\\u7ed3\\u6784\\uff0c\\u7ec4\\u4ef6\\u6d4b\\u8bd5\\u3001\\u96c6\\u6210\\u6d4b\\u8bd5\\u4e0e\\u590d\\u8dd1\\u8bc1\\u636e\\u9f50\\u5168\\u3002")
    )
    doc.add_paragraph(
        u("\\u0034\\u002e\\u0032 \\u8d1f\\u8d23\\u4efb AI \\u4f7f\\u7528\\uff1aAI \\u5efa\\u8bae\\u5747\\u7ecf\\u8fc7\\u672c\\u5730\\u9a8c\\u8bc1\\u540e\\u624d\\u91c7\\u7eb3\\uff0c\\u4e0d\\u5b58\\u5728\\u672a\\u9a8c\\u8bc1\\u76f4\\u63a5\\u63d0\\u4ea4\\u60c5\\u51b5\\u3002")
    )

    doc.add_heading(u("\\u4e94\\u3001\\u6700\\u7ec8\\u63d0\\u4ea4\\u72b6\\u6001\\uff08v11\\uff09"), level=1)
    final_lines = [
        u("\\u6700\\u7ec8\\u76ee\\u5f55\\uff1a") + rel(v11),
        u("\\u6700\\u5c0f\\u63d0\\u4ea4\\u5305\\uff083\\u6587\\u4ef6\\uff09\\uff1a") + rel(mini),
        u("\\u5b8c\\u6574\\u63d0\\u4ea4\\u5305\\uff1a") + rel(full),
        u("\\u8d85\\u4e25\\u683c\\u5ba1\\u8ba1\\u62a5\\u544a\\uff1a") + rel(v11 / "V11_ULTRA_STRICT_FINAL_AUDIT.md"),
    ]
    for line in final_lines:
        doc.add_paragraph(line, style="List Bullet")

    out_path = ai_dir / "AI_Log_Detailed_CN_v11.docx"
    doc.save(str(out_path))
    print("WROTE:", out_path)


if __name__ == "__main__":
    main()
